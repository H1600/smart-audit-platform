"""审计底稿生成器 —— 自动生成 Word/Excel 审计底稿

支持:
- 标准审计底稿（含摘要、明细、异常、结论）
- 风险导向底稿
- 科目专项底稿
- 一键导出 Word
"""
from __future__ import annotations

import json
import logging
from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from sqlalchemy.orm import Session

from ..models import LedgerRecord, AuditReport, TaskJob
from ..settings import EXPORT_DIR
from .tools import tool_get_account_summary, tool_detect_anomalies, tool_reconcile_account

logger = logging.getLogger(__name__)


def _add_heading(doc: Document, text: str, level: int = 1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x0B, 0x1A, 0x2E)
    return h


def _add_table(doc: Document, headers: list[str], rows: list[list[Any]], col_widths: list[float] | None = None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = str(h)
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)

    # Data
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val) if val is not None else ""
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)

    return table


def generate_workpaper_docx(db: Session, task_id: int, format: str = "standard") -> Path:
    """生成审计底稿 Word 文档"""
    task = db.get(TaskJob, task_id)
    if not task:
        raise ValueError("任务不存在")

    doc = Document()

    # 页面设置
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)

    # ── 封面 ──
    _add_heading(doc, "审计工作底稿", level=0)
    doc.add_paragraph(f"编制日期：{datetime.now().strftime('%Y年%m月%d日')}")
    doc.add_paragraph(f"被审计文件：{task.file.filename if task.file else '未知'}")
    doc.add_paragraph(f"处理状态：{task.status}")
    doc.add_paragraph("")

    # ── 一、审计摘要 ──
    _add_heading(doc, "一、审计摘要", level=1)

    records = db.query(LedgerRecord).filter(LedgerRecord.task_id == task_id).all()
    exceptions = [r for r in records if r.is_exception]
    total_debit = sum(r.debit or 0 for r in records)
    total_credit = sum(r.credit or 0 for r in records)

    summary_rows = [
        ["总记录数", str(len(records))],
        ["异常记录数", str(len(exceptions))],
        ["借方合计", f"¥{total_debit:,.2f}"],
        ["贷方合计", f"¥{total_credit:,.2f}"],
        ["借贷差额", f"¥{abs(total_debit - total_credit):,.2f}"],
        ["异常率", f"{len(exceptions)/max(len(records),1)*100:.1f}%"],
    ]
    _add_table(doc, ["指标", "数值"], summary_rows)
    doc.add_paragraph("")

    # ── 二、异常明细 ──
    _add_heading(doc, "二、异常明细", level=1)
    if exceptions:
        exception_rows = [
            [r.id, r.voucher_no or "", f"{r.account_code} {r.account_name}", r.summary[:40], f"{r.debit:,.2f}" if r.debit else "-", f"{r.credit:,.2f}" if r.credit else "-", r.exception_reason[:60]]
            for r in exceptions[:50]
        ]
        _add_table(doc, ["ID", "凭证号", "科目", "摘要", "借方", "贷方", "异常原因"], exception_rows)
    else:
        doc.add_paragraph("✅ 未发现异常记录")
    doc.add_paragraph("")

    # ── 三、科目汇总 ──
    _add_heading(doc, "三、科目汇总", level=1)
    by_account: dict[str, dict] = {}
    for r in records:
        code = r.account_code or "未映射"
        if code not in by_account:
            by_account[code] = {"name": r.account_name or "", "debit": 0.0, "credit": 0.0, "count": 0, "exc": 0}
        by_account[code]["debit"] += r.debit or 0
        by_account[code]["credit"] += r.credit or 0
        by_account[code]["count"] += 1
        if r.is_exception:
            by_account[code]["exc"] += 1

    account_rows = [
        [code, data["name"], data["count"], f"{data['debit']:,.2f}", f"{data['credit']:,.2f}", data["exc"]]
        for code, data in sorted(by_account.items(), key=lambda x: x[1]["debit"] + x[1]["credit"], reverse=True)
    ]
    _add_table(doc, ["科目编码", "科目名称", "笔数", "借方合计", "贷方合计", "异常数"], account_rows[:30])
    doc.add_paragraph("")

    # ── 四、勾稽校验结果 ──
    _add_heading(doc, "四、勾稽校验结果", level=1)
    reports = db.query(AuditReport).filter(AuditReport.task_id == task_id).all()
    if reports:
        report_rows = [[r.rule_name, "通过" if r.passed else "未通过", r.details[:100]] for r in reports]
        _add_table(doc, ["校验规则", "结果", "详情"], report_rows)
    else:
        doc.add_paragraph("暂无勾稽校验记录")
    doc.add_paragraph("")

    # ── 五、审计结论 ──
    _add_heading(doc, "五、审计结论", level=1)
    risk = "低" if len(exceptions) / max(len(records), 1) < 0.05 else "中" if len(exceptions) / max(len(records), 1) < 0.15 else "高"
    doc.add_paragraph(f"风险等级：{risk}")
    doc.add_paragraph(f"审计意见：基于以上分析，该批次数据共 {len(records)} 条记录，发现 {len(exceptions)} 条异常，异常率 {len(exceptions)/max(len(records),1)*100:.1f}%。")
    doc.add_paragraph("建议对异常记录逐条复核，重点关注大额交易和科目匹配情况。")

    # 保存
    export_path = Path(EXPORT_DIR) / f"audit_workpaper_{task_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    doc.save(str(export_path))
    logger.info("审计底稿已生成: %s", export_path)
    return export_path
