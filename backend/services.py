from __future__ import annotations

import csv
import html
import json
import logging
import os
import re
import shutil
import tempfile
import uuid
import importlib.util
import sys
from functools import lru_cache
from datetime import date, datetime
from pathlib import Path
from statistics import median
from typing import Any

from sqlalchemy import Select, select
from sqlalchemy.orm import Session

from .models import AuditReport, FileUpload, LedgerRecord, OcrResult, TaskJob
from .settings import (
    ALLOWED_EXTENSIONS,
    EXPORT_DIR,
    OCR_ENGINE,
    OCR_LANGS,
    OCR_MIN_CONF,
    OCR_MODEL_DIR,
    OCR_TEXTLINE,
    UPLOAD_DIR,
)


DATE_RE = re.compile(r"(20\d{2}|19\d{2})[-/.年](\d{1,2})[-/.月](\d{1,2})")
AMOUNT_RE = re.compile(r"[-+]?\d{1,3}(?:,\d{3})*(?:\.\d+)?|[-+]?\d+(?:\.\d+)?")
VOUCHER_RE = re.compile(r"^[\w\u4e00-\u9fff]{1,8}[-_]?\d{1,8}$")
SPLIT_RE = re.compile(r"[\t|]+|\s{2,}")
ACCOUNT_HINTS = {
    "现金": ("1001", "库存现金"),
    "银行": ("1002", "银行存款"),
    "应收": ("1122", "应收账款"),
    "应付": ("2202", "应付账款"),
    "收入": ("6001", "主营业务收入"),
    "成本": ("6401", "主营业务成本"),
    "费用": ("6602", "管理费用"),
}


def _normalize_key(value: str) -> str:
    return re.sub(r"\s+", "", str(value or "")).replace("_", "").replace(":", "").replace("：", "").lower()


STANDARD_FIELDS = ["日期", "凭证号", "科目编码", "科目名称", "摘要", "借方", "贷方", "余额"]
FIELD_ALIASES = {
    "日期": {"日期", "记账日期", "业务日期", "凭证日期", "date"},
    "凭证号": {"凭证号", "凭证编号", "voucher", "voucher_no"},
    "科目编码": {"科目编码", "科目代码", "account_code"},
    "科目名称": {"科目名称", "会计科目", "科目", "account_name"},
    "摘要": {"摘要", "说明", "业务描述", "remark", "description"},
    "借方": {"借方", "借方金额", "debit"},
    "贷方": {"贷方", "贷方金额", "credit"},
    "余额": {"余额", "balance"},
}
FIELD_ALIAS_LOOKUP = {
    _normalize_key(alias): canonical
    for canonical, aliases in FIELD_ALIASES.items()
    for alias in aliases
}
logger = logging.getLogger(__name__)

os.environ["FLAGS_use_mkldnn"] = "0"
os.environ["FLAGS_enable_mkldnn"] = "0"
os.environ["FLAGS_use_onednn"] = "0"
os.environ["FLAGS_enable_onednn"] = "0"
os.environ["FLAGS_enable_pir_api"] = "0"
os.environ["FLAGS_enable_pir_infer"] = "0"
os.environ["FLAGS_use_pir_api"] = "0"


def _ensure_torch_dlls() -> None:
    try:
        spec = importlib.util.find_spec("torch")
        if not spec or not spec.submodule_search_locations:
            return
        torch_dir = Path(spec.submodule_search_locations[0])
        lib_dir = torch_dir / "lib"
        if lib_dir.exists():
            os.add_dll_directory(str(lib_dir))
            os.environ["PATH"] = f"{lib_dir}{os.pathsep}{os.environ.get('PATH', '')}"
    except Exception:
        pass


def ocr_diagnostics() -> dict[str, Any]:
    def _spec_exists(name: str) -> bool:
        return importlib.util.find_spec(name) is not None

    torch_spec = importlib.util.find_spec("torch")
    torch_dir = None
    torch_lib = None
    shm_exists = False
    if torch_spec and torch_spec.submodule_search_locations:
        torch_dir = Path(torch_spec.submodule_search_locations[0])
        torch_lib = torch_dir / "lib"
        shm_exists = bool(torch_lib and (torch_lib / "shm.dll").exists())

    paddleocr_err = ""
    try:
        import paddleocr  # noqa: F401
    except Exception as exc:
        paddleocr_err = f"{type(exc).__name__}: {exc}"

    paddle_err = ""
    try:
        import paddle  # noqa: F401
    except Exception as exc:
        paddle_err = f"{type(exc).__name__}: {exc}"

    return {
        "python": sys.version.split(" ")[0],
        "ocr_engine": OCR_ENGINE,
        "ocr_langs": OCR_LANGS,
        "ocr_min_conf": OCR_MIN_CONF,
        "paddleocr_installed": _spec_exists("paddleocr"),
        "paddle_installed": _spec_exists("paddle"),
        "torch_installed": _spec_exists("torch"),
        "torch_lib": str(torch_lib) if torch_lib else "",
        "torch_shm_exists": shm_exists,
        "paddleocr_import_error": paddleocr_err,
        "paddle_import_error": paddle_err,
        "ocr_model_dir": str(OCR_MODEL_DIR) if OCR_MODEL_DIR else "",
    }


def init_db_seed(db: Session) -> None:
    rows = [
        ("库存现金", "1001", "库存现金"),
        ("银行存款", "1002", "银行存款"),
        ("应收账款", "1122", "应收账款"),
        ("应付账款", "2202", "应付账款"),
        ("主营业务收入", "6001", "主营业务收入"),
        ("主营业务成本", "6401", "主营业务成本"),
        ("管理费用", "6602", "管理费用"),
    ]
    from .models import AccountMapping

    if db.execute(select(AccountMapping)).first():
        return
    for raw, code, name in rows:
        db.add(AccountMapping(raw_account=raw, standard_code=code, standard_name=name))
    db.commit()


def save_upload(upload_file, db: Session) -> tuple[FileUpload, TaskJob]:
    suffix = Path(upload_file.filename or "").suffix.lower()
    if suffix not in ALLOWED_EXTENSIONS:
        raise ValueError(f"不支持的文件类型：{suffix or '未知'}")

    safe_name = Path(upload_file.filename or f"upload{suffix}").name
    storage_name = f"{datetime.utcnow().strftime('%Y%m%d%H%M%S')}_{uuid.uuid4().hex}{suffix}"
    storage_path = UPLOAD_DIR / storage_name
    with storage_path.open("wb") as buffer:
        shutil.copyfileobj(upload_file.file, buffer)

    file_row = FileUpload(
        filename=safe_name,
        file_type=suffix.lstrip("."),
        size=storage_path.stat().st_size,
        storage_path=str(storage_path),
    )
    db.add(file_row)
    db.flush()
    task = TaskJob(file_id=file_row.id)
    db.add(task)
    db.commit()
    db.refresh(file_row)
    db.refresh(task)
    return file_row, task


def process_task(task_id: int, session_factory) -> None:
    db: Session = session_factory()
    try:
        logger.info("task %s started", task_id)
        task = db.get(TaskJob, task_id)
        if not task:
            return
        file_row = db.get(FileUpload, task.file_id)
        if not file_row:
            raise RuntimeError("关联文件不存在")

        _mark(db, task, "running", "OCR/文件识别", 10, "开始读取上传文件")
        rows, ocr_pages = extract_source(Path(file_row.storage_path), file_row.file_type)
        for page in ocr_pages:
            db.add(OcrResult(task_id=task.id, **page))
        db.commit()

        _mark(db, task, "running", "ETL 清洗与标准化", 45, f"识别到 {len(rows)} 行候选数据")
        records = normalize_records(rows, task.id, file_row.id)
        for record in records:
            db.add(record)
        db.commit()

        _mark(db, task, "running", "勾稽关系校验", 75, "执行借贷平衡、金额和字段完整性校验")
        reports = build_reports(db, task.id)
        for report in reports:
            db.add(report)
        task.status = "completed"
        task.current_step = "已完成"
        task.progress = 100
        task.ended_at = datetime.utcnow()
        file_row.status = "completed"
        task.logs = _append_log(task.logs, "处理完成")
        db.commit()
        logger.info("task %s completed", task_id)
    except Exception as exc:
        logger.exception("task %s failed", task_id)
        task = db.get(TaskJob, task_id)
        if task:
            task.status = "failed"
            task.current_step = "失败"
            task.error_info = str(exc)
            task.ended_at = datetime.utcnow()
            task.logs = _append_log(task.logs, f"处理失败：{exc}")
            if task.file:
                task.file.status = "failed"
            db.commit()
    finally:
        db.close()


def extract_source(path: Path, file_type: str) -> tuple[list[dict[str, Any]], list[dict[str, Any]]]:
    suffix = f".{file_type.lower()}"
    if suffix in {".xlsx", ".xls", ".csv"}:
        return extract_table(path)
    if suffix == ".pdf":
        text = extract_pdf_text(path)
        rows = parse_tabular_text(text)
        return rows, [{"page_no": 1, "raw_text": text, "structured_fields": json.dumps({"source": "pdf_text", "parsed_rows": len(rows)}, ensure_ascii=False), "confidence": 0.86, "source_bbox": ""}]
    text, rows, confidence, structured = extract_image_text(path)
    return rows, [{"page_no": 1, "raw_text": text, "structured_fields": json.dumps(structured, ensure_ascii=False), "confidence": confidence, "source_bbox": ""}]


def extract_image_text(path: Path) -> tuple[str, list[dict[str, Any]], float, dict[str, Any]]:
    ocr = load_paddleocr()
    if ocr is None:
        text = f"图片文件 {path.name} 已接收。未安装 PaddleOCR 时使用离线占位识别，可在设置中配置 OCR 模型路径后替换。"
        rows = rows_from_text(text)
        return text, rows, 0.55, {"source": "image_placeholder", "reason": "paddleocr_not_installed"}

    ocr_path, cleanup = _prepare_ocr_image(path)
    try:
        result = ocr.ocr(ocr_path)
        items = _flatten_paddleocr_result(result)
        filtered = [item for item in items if item.get("score", 0.0) >= OCR_MIN_CONF and item.get("text")]
        lines = _ocr_items_to_lines(filtered or items)
        text = "\n".join(lines) if lines else f"图片文件 {path.name} 未识别到可用文本"
        rows = parse_tabular_text(text)
        confidence = round(sum(item["score"] for item in filtered) / len(filtered), 2) if filtered else 0.78
        structured = {
            "source": _ocr_engine_name(ocr),
            "items": len(items),
            "filtered_items": len(filtered),
            "lines": len(lines),
            "parsed_rows": len(rows),
            "min_conf": OCR_MIN_CONF,
        }
        return text, rows, confidence, structured
    except Exception as exc:
        text = f"图片文件 {path.name} OCR 失败：{exc}"
        rows = rows_from_text(text)
        return text, rows, 0.4, {"source": "image_error", "error": str(exc)}
    finally:
        if cleanup:
            try:
                os.remove(ocr_path)
            except OSError:
                pass


def extract_table(path: Path) -> tuple[list[dict[str, Any]], list[dict[str, Any]]]:
    rows: list[dict[str, Any]] = []
    try:
        matrix = _load_table_matrix(path)
        rows = _matrix_to_rows(matrix)
    except Exception:
        rows = rows_from_text(path.read_bytes()[:2048].decode("utf-8", errors="ignore"))
    raw = "\n".join(json.dumps(row, ensure_ascii=False) for row in rows[:30]) or "未识别到表格行"
    return rows, [{"page_no": 1, "raw_text": raw, "structured_fields": json.dumps({"source": "table"}, ensure_ascii=False), "confidence": 0.92, "source_bbox": ""}]


def extract_pdf_text(path: Path) -> str:
    try:
        from PyPDF2 import PdfReader

        reader = PdfReader(str(path))
        return "\n".join(page.extract_text() or "" for page in reader.pages).strip() or f"PDF {path.name} 未提取到文本"
    except Exception as exc:
        return f"PDF {path.name} 文本提取失败：{exc}"


def rows_from_text(text: str) -> list[dict[str, Any]]:
    parsed = parse_tabular_text(text)
    if parsed:
        return parsed
    rows = []
    for idx, line in enumerate([line.strip() for line in text.splitlines() if line.strip()], start=1):
        rows.append({"row_no": idx, "摘要": line})
    if not rows:
        rows.append({"row_no": 1, "摘要": "未识别到明细，生成待复核占位记录"})
    return rows


def parse_tabular_text(text: str) -> list[dict[str, Any]]:
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    if not lines:
        return [{"row_no": 1, "摘要": "未识别到明细，生成待复核占位记录"}]

    matrix = [_split_table_line(line) for line in lines]
    header_index = _detect_header_row(matrix)
    if header_index is not None:
        return _matrix_rows_with_headers(matrix, header_index)

    rows: list[dict[str, Any]] = []
    for idx, tokens in enumerate(matrix, start=1):
        rows.append(_row_from_tokens(tokens, idx))
    return rows


def _load_table_matrix(path: Path) -> list[list[str]]:
    suffix = path.suffix.lower()
    if suffix == ".csv":
        with path.open("r", encoding="utf-8-sig", newline="") as fh:
            matrix = [[str(cell).strip() for cell in row] for row in csv.reader(fh)]
        return [row for row in matrix if any(cell for cell in row)]

    if suffix == ".xlsx":
        from openpyxl import load_workbook

        workbook = load_workbook(path, data_only=True)
        worksheet = workbook.active
        matrix = [["" if cell is None else str(cell).strip() for cell in row] for row in worksheet.iter_rows(values_only=True)]
        for merged_range in worksheet.merged_cells.ranges:
            min_col, min_row, max_col, max_row = merged_range.bounds
            source = matrix[min_row - 1][min_col - 1] if min_row - 1 < len(matrix) and min_col - 1 < len(matrix[min_row - 1]) else ""
            for row_idx in range(min_row - 1, min(max_row, len(matrix))):
                for col_idx in range(min_col - 1, min(max_col, len(matrix[row_idx]))):
                    if not matrix[row_idx][col_idx]:
                        matrix[row_idx][col_idx] = source
        return [row for row in matrix if any(cell for cell in row)]

    try:
        import pandas as pd

        frame = pd.read_excel(path, header=None, dtype=str).fillna("")
        matrix = [[str(cell).strip() for cell in row] for row in frame.values.tolist()]
        return [row for row in matrix if any(cell for cell in row)]
    except Exception:
        with path.open("r", encoding="utf-8-sig", errors="ignore") as fh:
            return [[segment.strip() for segment in line.split(",")] for line in fh if line.strip()]


def _matrix_rows_with_headers(matrix: list[list[str]], header_index: int) -> list[dict[str, Any]]:
    header_row = matrix[header_index]
    headers = _normalize_headers(header_row)
    rows: list[dict[str, Any]] = []
    for idx, row in enumerate(matrix[header_index + 1 :], start=header_index + 2):
        if not any(cell.strip() for cell in row if cell is not None):
            continue
        row_dict: dict[str, Any] = {"row_no": idx}
        width = max(len(headers), len(row))
        normalized_headers = _expand_headers(headers, width)
        for col_idx in range(width):
            value = row[col_idx].strip() if col_idx < len(row) and row[col_idx] is not None else ""
            if not value:
                continue
            row_dict[normalized_headers[col_idx]] = value
        if len(row_dict) > 1:
            rows.append(row_dict)
    return rows or [_row_from_tokens(matrix[header_index], header_index + 1)]


def _normalize_headers(headers: list[str]) -> list[str]:
    normalized: list[str] = []
    seen: dict[str, int] = {}
    for idx, header in enumerate(headers):
        canonical = _canonical_field_name(header)
        if not canonical:
            canonical = f"列{idx + 1}"
        count = seen.get(canonical, 0)
        seen[canonical] = count + 1
        if count:
            canonical = f"{canonical}_{count + 1}"
        normalized.append(canonical)
    return normalized


def _expand_headers(headers: list[str], width: int) -> list[str]:
    if len(headers) >= width:
        return headers[:width]
    expanded = list(headers)
    for idx in range(len(headers), width):
        expanded.append(f"列{idx + 1}")
    return expanded


def _matrix_to_rows(matrix: list[list[str]]) -> list[dict[str, Any]]:
    cleaned = [["" if cell is None else str(cell).strip() for cell in row] for row in matrix if any((cell or "").strip() for cell in row)]
    if not cleaned:
        return [{"row_no": 1, "摘要": "未识别到表格行"}]

    header_index = _detect_header_row(cleaned)
    if header_index is not None:
        return _matrix_rows_with_headers(cleaned, header_index)

    rows: list[dict[str, Any]] = []
    for idx, row in enumerate(cleaned, start=1):
        rows.append(_row_from_tokens(row, idx))
    return rows


def _detect_header_row(matrix: list[list[str]]) -> int | None:
    best_index: int | None = None
    best_score = 0
    for index, row in enumerate(matrix[:8]):
        score = sum(1 for cell in row if _canonical_field_name(cell) in STANDARD_FIELDS)
        if score > best_score and score >= 2:
            best_index = index
            best_score = score
    return best_index


def _row_from_tokens(tokens: list[str], row_no: int) -> dict[str, Any]:
    cells = [str(token).strip() for token in tokens if str(token).strip()]
    if not cells:
        return {"row_no": row_no, "摘要": "未识别到明细，生成待复核占位记录"}

    row: dict[str, Any] = {"row_no": row_no}
    used: set[int] = set()

    for index, token in enumerate(cells):
        parsed_date = _parse_date(token)
        if parsed_date and "日期" not in row:
            row["日期"] = parsed_date.isoformat()
            used.add(index)
            continue
        if not row.get("凭证号") and _looks_like_voucher(token):
            row["凭证号"] = token
            used.add(index)
            continue
        if not row.get("科目编码") and re.fullmatch(r"\d{3,6}", token):
            row["科目编码"] = token
            used.add(index)
            continue
        if row.get("科目编码") and not row.get("科目名称") and not _is_amount_like(token) and not _looks_like_voucher(token):
            row["科目名称"] = token
            used.add(index)
            continue

    amounts = [token for token in cells if _is_amount_like(token)]
    if len(amounts) >= 3:
        row["借方"], row["贷方"], row["余额"] = amounts[-3:]
    elif len(amounts) == 2:
        row["借方"], row["贷方"] = amounts[-2:]
    elif len(amounts) == 1:
        row["借方"] = amounts[0]

    summary_parts = [token for index, token in enumerate(cells) if index not in used and not _is_amount_like(token)]
    if summary_parts:
        row["摘要"] = " ".join(summary_parts)
    else:
        row["摘要"] = "待复核"

    if "科目名称" not in row and row.get("摘要"):
        guessed_account = _guess_account(row["摘要"])
        row["科目编码"] = row.get("科目编码") or guessed_account[0]
        row["科目名称"] = row.get("科目名称") or guessed_account[1]
    return row


def _split_table_line(line: str) -> list[str]:
    if not line:
        return []
    if "|" in line or "\t" in line or re.search(r"\s{2,}", line):
        tokens = [piece.strip() for piece in SPLIT_RE.split(line) if piece.strip()]
        if tokens:
            return tokens
    if line.count(",") >= 2 and "，" not in line:
        tokens = [piece.strip() for piece in line.split(",") if piece.strip()]
        if tokens:
            return tokens
    return [piece.strip() for piece in re.split(r"\s+", line) if piece.strip()]


def _is_amount_like(value: str) -> bool:
    normalized = str(value).replace(",", "").strip()
    return bool(normalized) and bool(AMOUNT_RE.fullmatch(normalized))


def _looks_like_voucher(value: str) -> bool:
    return bool(VOUCHER_RE.fullmatch(str(value).strip()))


def _canonical_field_name(value: str) -> str:
    normalized = _normalize_key(value)
    return FIELD_ALIAS_LOOKUP.get(normalized, str(value).strip())


def _parse_ocr_langs() -> tuple[str, list[str]]:
    raw = [lang.strip().lower() for lang in (OCR_LANGS or []) if lang.strip()]
    if not raw:
        return "ch", ["ch_sim", "en"]
    paddle_lang = "ch" if any(lang in {"zh", "zh-cn", "ch", "cn", "ch_sim"} for lang in raw) else "en"
    easy_langs: list[str] = []
    for lang in raw:
        if lang in {"zh", "zh-cn", "ch", "cn", "ch_sim", "zh_cn"}:
            easy_langs.append("ch_sim")
        elif lang in {"en", "english"}:
            easy_langs.append("en")
        else:
            easy_langs.append(lang)
    if "en" not in easy_langs:
        easy_langs.append("en")
    return paddle_lang, easy_langs


def _prepare_ocr_image(path: Path) -> tuple[str, bool]:
    try:
        from PIL import Image, ImageEnhance, ImageOps
    except Exception:
        return str(path), False

    try:
        img = Image.open(path)
        img = ImageOps.exif_transpose(img)
        img = img.convert("L")
        img = ImageOps.autocontrast(img)
        img = ImageEnhance.Sharpness(img).enhance(1.4)
        img = img.convert("RGB")
        suffix = path.suffix if path.suffix else ".png"
        handle = tempfile.NamedTemporaryFile(prefix="ocr_", suffix=suffix, dir=str(path.parent), delete=False)
        img.save(handle.name)
        handle.close()
        return handle.name, True
    except Exception:
        return str(path), False


@lru_cache(maxsize=1)
def load_paddleocr():
    """Try PaddleOCR first, then EasyOCR; return None if both unavailable."""
    engine = (OCR_ENGINE or "auto").strip().lower()
    paddle_lang, easy_langs = _parse_ocr_langs()
    if engine in {"off", "none", "disabled"}:
        return None

    _ensure_torch_dlls()

    if engine in {"auto", "paddle", "paddleocr"}:
        try:
            from paddleocr import PaddleOCR

            kwargs = {"use_textline_orientation": OCR_TEXTLINE, "lang": paddle_lang}
            if OCR_MODEL_DIR is not None:
                det_dir = OCR_MODEL_DIR / "det"
                rec_dir = OCR_MODEL_DIR / "rec"
                cls_dir = OCR_MODEL_DIR / "cls"
                if det_dir.exists():
                    kwargs["det_model_dir"] = str(det_dir)
                if rec_dir.exists():
                    kwargs["rec_model_dir"] = str(rec_dir)
                if cls_dir.exists():
                    kwargs["cls_model_dir"] = str(cls_dir)

            ocr = PaddleOCR(**kwargs)
            ocr._engine_name = "paddleocr"
            return ocr
        except Exception:
            if engine not in {"auto"}:
                return None

    if engine in {"auto", "easy", "easyocr"}:
        try:
            import easyocr

            reader = easyocr.Reader(easy_langs, gpu=False)
            return _EasyOCRFallback(reader)
        except Exception:
            return None
    return None


class _EasyOCRFallback:
    """Adapter that wraps EasyOCR results into the same interface as PaddleOCR."""

    def __init__(self, reader) -> None:
        self._reader = reader
        self._engine_name = "easyocr"

    def ocr(self, img_path: str) -> list[list[tuple[list, tuple[str, float]]]]:
        result = self._reader.readtext(img_path)
        adapted: list[list[tuple[list, tuple[str, float]]]] = []
        for bbox, text, score in result:
            adapted.append([([int(x) for coords in bbox for x in coords], (text, score))])
        return adapted


def _flatten_paddleocr_result(result: Any) -> list[dict[str, Any]]:
    if not result:
        return []

    items = result
    if isinstance(result, list) and len(result) == 1 and isinstance(result[0], list):
        items = result[0]

    flattened: list[dict[str, Any]] = []
    for item in items:
        text = ""
        score = 0.0
        bbox = ""
        rect: tuple[float, float, float, float] | None = None
        if isinstance(item, (list, tuple)) and len(item) == 2:
            rect = _bbox_to_rect(item[0])
            bbox = json.dumps(item[0], ensure_ascii=False)
            if isinstance(item[1], (list, tuple)) and item[1]:
                text = str(item[1][0])
                score = float(item[1][1]) if len(item[1]) > 1 else 0.0
            else:
                text = str(item[1])
        elif isinstance(item, str):
            text = item
        if text:
            payload: dict[str, Any] = {"text": text, "score": score, "bbox": bbox}
            if rect:
                payload.update(
                    {
                        "x_min": rect[0],
                        "y_min": rect[1],
                        "x_max": rect[2],
                        "y_max": rect[3],
                        "width": rect[2] - rect[0],
                        "height": rect[3] - rect[1],
                    }
                )
            flattened.append(payload)
    return flattened


def _ocr_engine_name(ocr) -> str:
    return getattr(ocr, "_engine_name", ocr.__class__.__name__.lower())


def _bbox_to_rect(bbox: Any) -> tuple[float, float, float, float] | None:
    if bbox is None:
        return None
    if isinstance(bbox, (list, tuple)):
        if bbox and isinstance(bbox[0], (list, tuple)):
            xs = [float(point[0]) for point in bbox if len(point) >= 2]
            ys = [float(point[1]) for point in bbox if len(point) >= 2]
        else:
            coords = [float(value) for value in bbox if isinstance(value, (int, float))]
            if len(coords) < 4:
                return None
            xs = coords[0::2]
            ys = coords[1::2]
        if not xs or not ys:
            return None
        return min(xs), min(ys), max(xs), max(ys)
    return None


def _ocr_items_to_lines(items: list[dict[str, Any]]) -> list[str]:
    if not items:
        return []
    grouped = _group_ocr_items(items)
    return [line for line in (_line_from_items(group) for group in grouped) if line]


def _group_ocr_items(items: list[dict[str, Any]]) -> list[list[dict[str, Any]]]:
    with_box = [item for item in items if item.get("y_min") is not None]
    if not with_box:
        return [[item] for item in items]

    heights = [item.get("height") or 0 for item in with_box if item.get("height")]
    base_height = median(heights) if heights else 12
    threshold = max(10.0, base_height * 0.6)
    sorted_items = sorted(with_box, key=lambda item: (item.get("y_min", 0), item.get("x_min", 0)))
    lines: list[dict[str, Any]] = []
    for item in sorted_items:
        center = ((item.get("y_min", 0) + item.get("y_max", 0)) / 2.0) if item.get("y_min") is not None else 0.0
        placed = False
        for line in lines:
            if abs(center - line["center"]) <= threshold:
                line["items"].append(item)
                count = line["count"] + 1
                line["center"] = (line["center"] * line["count"] + center) / count
                line["count"] = count
                placed = True
                break
        if not placed:
            lines.append({"center": center, "count": 1, "items": [item]})

    return [line["items"] for line in sorted(lines, key=lambda line: line["center"]) if line["items"]]


def _line_from_items(items: list[dict[str, Any]]) -> str:
    ordered = sorted(items, key=lambda item: item.get("x_min", 0))
    gaps = []
    for prev, curr in zip(ordered, ordered[1:]):
        gaps.append(max(0.0, float(curr.get("x_min", 0)) - float(prev.get("x_max", 0))))
    base_gap = median(gaps) if gaps else 0.0
    gap_threshold = max(15.0, base_gap * 1.5)

    parts: list[str] = []
    prev = None
    for item in ordered:
        if prev is not None:
            gap = float(item.get("x_min", 0)) - float(prev.get("x_max", 0))
            parts.append("  " if gap > gap_threshold else " ")
        parts.append(str(item.get("text", "")))
        prev = item
    return "".join(parts).strip()


def normalize_records(rows: list[dict[str, Any]], task_id: int, file_id: int) -> list[LedgerRecord]:
    # Try loading ML classifier for better account prediction
    try:
        from .account_classifier import load_classifier, _tokenize, _keyword_fallback

        clf_pair = load_classifier()
    except Exception:
        clf_pair = None

    seen: set[tuple[Any, ...]] = set()
    records: list[LedgerRecord] = []
    for row in rows:
        compact = {str(k).strip(): "" if v is None else str(v).strip() for k, v in row.items()}
        summary = _pick(compact, ["摘要", "说明", "业务描述", "remark", "description"]) or "待复核"
        record_date = _parse_date(_pick(compact, ["日期", "记账日期", "业务日期", "date"]) or summary)
        voucher_no = _pick(compact, ["凭证号", "凭证编号", "voucher", "voucher_no"]) or _guess_voucher(summary)
        account_code = _pick(compact, ["科目编码", "科目代码", "account_code"]) or ""
        account_name = _pick(compact, ["科目名称", "会计科目", "科目", "account_name"]) or ""

        # ML prediction fallback when account info is missing
        if (not account_code or not account_name or account_name == "未映射") and summary and summary != "待复核":
            try:
                if clf_pair is not None:
                    vect, clf = clf_pair
                    tokens_list = [_tokenize(summary)]
                    X = vect.transform(tokens_list)
                    probs = clf.predict_proba(X)[0]
                    best_idx = int(probs.argmax())
                    pred_name = str(clf.classes_[best_idx])
                    pred_conf = round(float(probs[best_idx]), 4)
                    if pred_conf >= 0.3:
                        if not account_name or account_name == "未映射":
                            account_name = pred_name
            except Exception:
                pass

        # Keyword fallback
        if not account_code:
            kw_code, kw_name = _keyword_fallback(account_name or summary)
            account_code = account_code or kw_code
            if not account_name or account_name == "未映射":
                account_name = kw_name

        # Also try old _guess_account for backward compat
        if not account_code:
            guessed_code, guessed_name = _guess_account(account_name or summary)
            account_code = account_code or guessed_code
            if not account_name or account_name == "未映射":
                account_name = guessed_name

        debit = _parse_amount(_pick(compact, ["借方", "借方金额", "debit"]))
        credit = _parse_amount(_pick(compact, ["贷方", "贷方金额", "credit"]))
        balance = _parse_amount(_pick(compact, ["余额", "balance"]))
        if debit == 0 and credit == 0:
            amount = _first_amount(summary)
            debit = amount if amount >= 0 else 0
            credit = abs(amount) if amount < 0 else 0

        dedupe_key = (record_date, voucher_no, account_code, account_name, summary, debit, credit)
        if dedupe_key in seen:
            continue
        seen.add(dedupe_key)

        reasons = []
        if not record_date:
            reasons.append("日期缺失或格式异常")
        if not account_code or not account_name:
            reasons.append("科目信息不完整")
        if debit < 0 or credit < 0:
            reasons.append("借贷金额不应为负")
        if debit and credit:
            reasons.append("同一行同时存在借方和贷方金额")
        if not debit and not credit:
            reasons.append("借贷金额均为空")

        records.append(
            LedgerRecord(
                task_id=task_id,
                file_id=file_id,
                record_date=record_date,
                voucher_no=voucher_no,
                account_code=account_code,
                account_name=account_name,
                summary=summary[:500],
                debit=debit,
                credit=credit,
                balance=balance,
                source_page=1,
                source_row=int(compact.get("row_no") or 0),
                source_text=json.dumps(compact, ensure_ascii=False),
                is_exception=bool(reasons),
                exception_reason="；".join(reasons),
            )
        )
    return records


def build_reports(db: Session, task_id: int) -> list[AuditReport]:
    records = list(db.scalars(select(LedgerRecord).where(LedgerRecord.task_id == task_id)))
    debit_total = round(sum(r.debit for r in records), 2)
    credit_total = round(sum(r.credit for r in records), 2)
    exceptions = [r for r in records if r.is_exception]
    return [
        AuditReport(
            task_id=task_id,
            rule_name="借贷发生额平衡",
            passed=debit_total == credit_total,
            details=json.dumps({"debit_total": debit_total, "credit_total": credit_total, "difference": round(debit_total - credit_total, 2)}, ensure_ascii=False),
        ),
        AuditReport(
            task_id=task_id,
            rule_name="字段完整性与金额有效性",
            passed=not exceptions,
            details=json.dumps([{"record_id": r.id, "reason": r.exception_reason} for r in exceptions], ensure_ascii=False),
        ),
    ]


def query_records(
    db: Session,
    start_date: date | None = None,
    end_date: date | None = None,
    account: str | None = None,
    voucher_no: str | None = None,
    min_amount: float | None = None,
    max_amount: float | None = None,
    page: int = 1,
    page_size: int = 20,
) -> dict[str, Any]:
    stmt: Select = select(LedgerRecord)
    if start_date:
        stmt = stmt.where(LedgerRecord.record_date >= start_date)
    if end_date:
        stmt = stmt.where(LedgerRecord.record_date <= end_date)
    if account:
        like = f"%{account}%"
        stmt = stmt.where((LedgerRecord.account_code.like(like)) | (LedgerRecord.account_name.like(like)))
    if voucher_no:
        stmt = stmt.where(LedgerRecord.voucher_no.like(f"%{voucher_no}%"))
    records = list(db.scalars(stmt.order_by(LedgerRecord.id.desc())))
    if min_amount is not None:
        records = [r for r in records if max(r.debit, r.credit, abs(r.balance)) >= min_amount]
    if max_amount is not None:
        records = [r for r in records if max(r.debit, r.credit, abs(r.balance)) <= max_amount]
    total = len(records)
    offset = max(page - 1, 0) * page_size
    return {"total": total, "page": page, "page_size": page_size, "items": [record_to_dict(r) for r in records[offset : offset + page_size]]}


def record_to_dict(record: LedgerRecord) -> dict[str, Any]:
    return {
        "id": record.id,
        "task_id": record.task_id,
        "file_id": record.file_id,
        "date": record.record_date.isoformat() if record.record_date else "",
        "voucher_no": record.voucher_no,
        "account_code": record.account_code,
        "account_name": record.account_name,
        "summary": record.summary,
        "debit": record.debit,
        "credit": record.credit,
        "balance": record.balance,
        "source_page": record.source_page,
        "source_row": record.source_row,
        "source_text": record.source_text,
        "is_exception": record.is_exception,
        "exception_reason": record.exception_reason,
    }


def export_task(db: Session, task_id: int, fmt: str) -> Path:
    records = list(db.scalars(select(LedgerRecord).where(LedgerRecord.task_id == task_id).order_by(LedgerRecord.id)))
    if fmt == "xbrl":
        path = EXPORT_DIR / f"task_{task_id}_xbrl.xml"
        body = ["<?xml version=\"1.0\" encoding=\"UTF-8\"?>", "<AuditLedger>"]
        for record in records:
            body.append(
                "  <LedgerRecord "
                f"id=\"{record.id}\" date=\"{record.record_date or ''}\" voucher=\"{html.escape(record.voucher_no)}\" "
                f"accountCode=\"{html.escape(record.account_code)}\" accountName=\"{html.escape(record.account_name)}\" "
                f"debit=\"{record.debit:.2f}\" credit=\"{record.credit:.2f}\" balance=\"{record.balance:.2f}\">"
                f"{html.escape(record.summary)}</LedgerRecord>"
            )
        body.append("</AuditLedger>")
        path.write_text("\n".join(body), encoding="utf-8")
        return path
    if fmt == "report":
        path = EXPORT_DIR / f"task_{task_id}_report.json"
        reports = list(db.scalars(select(AuditReport).where(AuditReport.task_id == task_id)))
        path.write_text(json.dumps([report_to_dict(r) for r in reports], ensure_ascii=False, indent=2), encoding="utf-8")
        return path
    if fmt == "docx":
        return _export_docx(db, task_id, records)

    path = EXPORT_DIR / f"task_{task_id}_records.xlsx"
    try:
        from openpyxl import Workbook

        wb = Workbook()
        ws = wb.active
        ws.title = "审计明细"
        headers = ["日期", "凭证号", "科目编码", "科目名称", "摘要", "借方", "贷方", "余额", "异常", "异常原因"]
        ws.append(headers)
        for r in records:
            ws.append([r.record_date, r.voucher_no, r.account_code, r.account_name, r.summary, r.debit, r.credit, r.balance, "是" if r.is_exception else "否", r.exception_reason])
        wb.save(path)
    except Exception:
        path = EXPORT_DIR / f"task_{task_id}_records.csv"
        with path.open("w", encoding="utf-8-sig", newline="") as fh:
            writer = csv.writer(fh)
            writer.writerow(["日期", "凭证号", "科目编码", "科目名称", "摘要", "借方", "贷方", "余额", "异常", "异常原因"])
            for r in records:
                writer.writerow([r.record_date, r.voucher_no, r.account_code, r.account_name, r.summary, r.debit, r.credit, r.balance, r.is_exception, r.exception_reason])
    return path


def report_to_dict(report: AuditReport) -> dict[str, Any]:
    return {
        "id": report.id,
        "task_id": report.task_id,
        "rule_name": report.rule_name,
        "passed": report.passed,
        "details": json.loads(report.details) if report.details else "",
        "generated_at": report.generated_at.isoformat(),
    }


def task_to_dict(task: TaskJob) -> dict[str, Any]:
    return {
        "id": task.id,
        "file_id": task.file_id,
        "filename": task.file.filename if task.file else "",
        "status": task.status,
        "current_step": task.current_step,
        "progress": task.progress,
        "started_at": task.started_at.isoformat() if task.started_at else "",
        "ended_at": task.ended_at.isoformat() if task.ended_at else "",
        "error_info": task.error_info or "",
        "logs": task.logs.splitlines()[-80:],
    }


def _mark(db: Session, task: TaskJob, status: str, step: str, progress: int, log: str) -> None:
    task.status = status
    task.current_step = step
    task.progress = progress
    if not task.started_at:
        task.started_at = datetime.utcnow()
    task.logs = _append_log(task.logs, log)
    if task.file:
        task.file.status = status
    db.commit()


def _append_log(existing: str, message: str) -> str:
    return f"{existing}{datetime.utcnow().isoformat(timespec='seconds')} {message}\n"


def _pick(row: dict[str, str], names: list[str]) -> str:
    normalized = {k.lower().replace(" ", "").replace("_", ""): v for k, v in row.items()}
    for name in names:
        key = name.lower().replace(" ", "").replace("_", "")
        if normalized.get(key):
            return normalized[key]
    return ""


def _parse_date(value: str) -> date | None:
    match = DATE_RE.search(value or "")
    if not match:
        return None
    y, m, d = [int(x) for x in match.groups()]
    try:
        return date(y, m, d)
    except ValueError:
        return None


def _parse_amount(value: str | None) -> float:
    if not value:
        return 0.0
    match = AMOUNT_RE.search(str(value).replace(",", ""))
    return round(float(match.group(0)), 2) if match else 0.0


def _first_amount(text: str) -> float:
    numbers = [_parse_amount(item) for item in AMOUNT_RE.findall(text or "")]
    return numbers[-1] if numbers else 0.0


def _guess_voucher(text: str) -> str:
    match = re.search(r"(记|收|付|转)?\s?字?\s?第?\s?([A-Za-z0-9-]{2,})\s?号?", text or "")
    return match.group(0).strip() if match else ""


def _guess_account(text: str) -> tuple[str, str]:
    for hint, account in ACCOUNT_HINTS.items():
        if hint in (text or ""):
            return account
    return "", ""


def _export_docx(db: Session, task_id: int, records: list[LedgerRecord]) -> Path:
    """Generate a Word audit report with summary, exceptions, and record details."""
    try:
        from docx import Document
        from docx.shared import Inches, Pt, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml.ns import qn
    except ImportError:
        raise RuntimeError("python-docx 未安装，请执行 pip install python-docx")

    reports = list(db.scalars(select(AuditReport).where(AuditReport.task_id == task_id)))
    exceptions = [r for r in records if r.is_exception]
    debit_total = round(sum(r.debit for r in records), 2)
    credit_total = round(sum(r.credit for r in records), 2)

    doc = Document()

    # ── Style setup ──
    style = doc.styles["Normal"]
    style.font.name = "微软雅黑"
    style.font.size = Pt(10.5)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

    # ── Title ──
    title = doc.add_heading("审计数据处理报告", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ── Meta ──
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_run = meta.add_run(f"任务编号: {task_id}  |  生成时间: {datetime.utcnow().strftime('%Y-%m-%d %H:%M')}  |  记录数: {len(records)}")
    meta_run.font.size = Pt(9)
    meta_run.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_paragraph("")  # spacer

    # ── Summary ──
    doc.add_heading("一、汇总信息", level=1)
    summary_data = [
        ("总记录数", str(len(records))),
        ("借方合计", f"¥{debit_total:,.2f}"),
        ("贷方合计", f"¥{credit_total:,.2f}"),
        ("借贷差额", f"¥{abs(debit_total - credit_total):,.2f}"),
        ("借贷平衡", "✅ 是" if debit_total == credit_total else "❌ 否"),
        ("异常记录数", str(len(exceptions))),
        ("正常记录数", str(len(records) - len(exceptions))),
    ]
    table = doc.add_table(rows=len(summary_data), cols=2, style="Light Grid Accent 1")
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (label, value) in enumerate(summary_data):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = value
        for cell in table.rows[i].cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)

    doc.add_paragraph("")

    # ── Validation Reports ──
    if reports:
        doc.add_heading("二、勾稽校验结果", level=1)
        for report in reports:
            status_text = "✅ 通过" if report.passed else "❌ 未通过"
            doc.add_paragraph(f"{status_text} — {report.rule_name}", style="List Bullet")
            if report.details:
                try:
                    detail = json.loads(report.details)
                    if isinstance(detail, dict):
                        for k, v in detail.items():
                            doc.add_paragraph(f"  {k}: {v}", style="List Bullet 2")
                except Exception:
                    doc.add_paragraph(f"  {report.details}")

    doc.add_paragraph("")

    # ── Exception Records ──
    if exceptions:
        doc.add_heading("三、异常记录明细", level=1)
        exc_table = doc.add_table(rows=1, cols=6, style="Light Grid Accent 1")
        exc_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr = exc_table.rows[0].cells
        headers = ["序号", "日期", "凭证号", "科目名称", "金额(借/贷)", "异常原因"]
        for i, text in enumerate(headers):
            hdr[i].text = text
            for p in hdr[i].paragraphs:
                for run in p.runs:
                    run.bold = True
                    run.font.size = Pt(9)

        for idx, r in enumerate(exceptions[:100], 1):
            row = exc_table.add_row()
            cells = row.cells
            cells[0].text = str(idx)
            cells[1].text = r.record_date.isoformat() if r.record_date else ""
            cells[2].text = r.voucher_no
            cells[3].text = r.account_name
            cells[4].text = f"{r.debit:.2f} / {r.credit:.2f}"
            cells[5].text = r.exception_reason
            for cell in cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.size = Pt(9)

    doc.add_paragraph("")

    # ── Full Record Table ──
    doc.add_heading("四、全部记录明细", level=1)
    rec_table = doc.add_table(rows=1, cols=8, style="Light Grid Accent 1")
    rec_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr2 = rec_table.rows[0].cells
    rec_headers = ["序号", "日期", "凭证号", "科目编码", "科目名称", "摘要", "借方", "贷方"]
    for i, text in enumerate(rec_headers):
        hdr2[i].text = text
        for p in hdr2[i].paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(8)

    for idx, r in enumerate(records[:500], 1):
        row = rec_table.add_row()
        cells = row.cells
        cells[0].text = str(idx)
        cells[1].text = r.record_date.isoformat() if r.record_date else ""
        cells[2].text = r.voucher_no
        cells[3].text = r.account_code
        cells[4].text = r.account_name
        cells[5].text = (r.summary or "")[:60]
        cells[6].text = f"{r.debit:.2f}"
        cells[7].text = f"{r.credit:.2f}"
        for cell in cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(8)

    # ── Footer ──
    doc.add_paragraph("")
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_run = footer.add_run("— 由智能财务审计数据处理平台自动生成 —")
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = RGBColor(150, 150, 150)
    footer_run.italic = True

    path = EXPORT_DIR / f"task_{task_id}_audit_report.docx"
    doc.save(str(path))
    return path
